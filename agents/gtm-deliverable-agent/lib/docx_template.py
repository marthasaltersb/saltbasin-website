from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x1B, 0x2A, 0x3B)
CREAM = RGBColor(0xF5, 0xF0, 0xE8)
GOLD = RGBColor(0xC4, 0x84, 0x3A)
RISK_RED = RGBColor(0xB5, 0x43, 0x3A)


def _shade_cell(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _section_header(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.color.rgb = GOLD
    run.font.bold = True
    run.font.size = Pt(13)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "C4843A")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    return p


def _add_table(doc: Document, headers: list[str], rows: list[list]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = CREAM
        _shade_cell(header_cells[i], "1B2A3B")
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = str(value) if value is not None else ""
    return table


def build_docx(deliverable: dict, output_path: Path) -> None:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Salt Basin Net Works")
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = NAVY
    title_run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(deliverable["topic"])
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = GOLD

    client_name = deliverable.get("engagement_client_name")
    if client_name:
        client_p = doc.add_paragraph()
        client_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        client_p.add_run(f"Prepared for {client_name}")

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_run = banner.add_run("DRAFT — HUMAN REVIEW REQUIRED BEFORE CLIENT USE")
    banner_run.font.color.rgb = RISK_RED
    banner_run.font.bold = True
    banner_run.italic = True

    doc.add_page_break()

    _section_header(doc, "Executive Summary")
    doc.add_paragraph(deliverable["executive_summary"])

    _section_header(doc, "Benchmark Master")
    _add_table(
        doc,
        ["Metric", "Value", "Source", "Year", "Relevance"],
        [
            [b["metric"], b["value"], b["source"], b["year"], b["relevance_note"]]
            for b in deliverable["benchmark_master"]
        ],
    )

    industry = deliverable.get("industry_breakdown") or []
    if industry:
        _section_header(doc, "Benchmark by Industry")
        _add_table(
            doc,
            ["Industry", "Mechanism", "Rate Estimate", "Source"],
            [
                [i["industry"], i["leakage_or_risk_mechanism"], i["rate_estimate"], i["rate_source"]]
                for i in industry
            ],
        )

    _section_header(doc, "Assumptions & Methodology")
    am = deliverable["assumptions_methodology"]

    doc.add_paragraph("Verified Primary-Source Statistics").runs[0].font.bold = True
    _add_table(
        doc,
        ["Statistic", "Value Used", "Primary Source", "Publication Date"],
        [
            [s["statistic_name"], s["value_used"], s["primary_source"], s["publication_date"]]
            for s in am["verified_statistics"]
        ],
    )

    doc.add_paragraph("Modeled Assumptions (Disclose in Presentations)").runs[0].font.bold = True
    _add_table(
        doc,
        ["Assumption", "Conservative", "Base", "Optimistic", "Recommendation"],
        [
            [m["assumption_name"], m["conservative"], m["base"], m["optimistic"], m["recommendation"]]
            for m in am["modeled_assumptions"]
        ],
    )

    doc.add_paragraph("Scenario-to-Source Mapping").runs[0].font.bold = True
    _add_table(
        doc,
        ["Scenario", "Mapped Source Category", "Confidence", "Note"],
        [
            [s["scenario"], s["mapped_source_category"], s["confidence_level"], s["note"]]
            for s in am["scenario_source_mapping"]
        ],
    )

    mapping = deliverable.get("client_mapping")
    if mapping:
        _section_header(doc, "Client Actuals vs. Benchmark")
        _add_table(
            doc,
            ["Metric", "Client Value", "Benchmark Value", "Delta", "Confidence"],
            [
                [
                    c["metric"],
                    c["client_value"] or "—",
                    c["benchmark_value"],
                    c["delta_description"],
                    c["confidence_level"],
                ]
                for c in mapping["client_actuals_vs_benchmark"]
            ],
        )

    gaps = deliverable.get("data_quality_gaps") or []
    if gaps:
        _section_header(doc, "Data Quality Gaps")
        _add_table(doc, ["Description", "Severity"], [[g["description"], g["severity"]] for g in gaps])

    flags = deliverable.get("unverified_flags") or []
    if flags:
        _section_header(doc, "Unverified Claims — Excluded")
        _add_table(
            doc, ["Claim", "Reason Unverified"], [[f["claim"], f["reason_unverified"]] for f in flags]
        )

    footer_paragraph = doc.sections[0].footer.paragraphs[0]
    footer_paragraph.text = (
        "Salt Basin Net Works | Bottom Lines with a Rising Tide | "
        "Sources verified per Assumptions & Methodology tab"
    )
    for run in footer_paragraph.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = NAVY

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
