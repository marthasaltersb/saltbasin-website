import json
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


SOURCES = [
    Path(r"C:\Users\mbets\Downloads\Universal_Scientific_Repository_Fully_Normalized_v3.0.xlsx"),
    Path(r"C:\Users\mbets\Downloads\Enterprise_Operating_System_Scientific_Equivalent_Repository_v0.1.xlsx"),
    Path(r"C:\Users\mbets\Downloads\Salt_Basin_Scientific_to_Enterprise_Mapping_Control_Layer_v0.2.xlsx"),
]
DOC = Path(r"C:\Users\mbets\Downloads\Salt_Basin_Genesis_Program_and_Technical_Specification_v1.0.docx")
OUT = Path(__file__).parent / "extracted"


def clean(value):
    if value is None:
        return None
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return value


def extract_workbook(path):
    wb = load_workbook(path, data_only=False, read_only=True)
    result = {"file": path.name, "sheets": []}
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows():
            values = [clean(cell.value) for cell in row]
            while values and values[-1] is None:
                values.pop()
            if any(value is not None for value in values):
                rows.append(values)
        result["sheets"].append({
            "name": ws.title,
            "max_row": ws.max_row,
            "max_column": ws.max_column,
            "rows": rows,
        })
    return result


def extract_doc(path):
    doc = Document(path)
    paragraphs = [
        {"style": p.style.name if p.style else None, "text": p.text.strip()}
        for p in doc.paragraphs if p.text.strip()
    ]
    tables = []
    for table in doc.tables:
        tables.append([[cell.text.strip() for cell in row.cells] for row in table.rows])
    return {"file": path.name, "paragraphs": paragraphs, "tables": tables}


OUT.mkdir(parents=True, exist_ok=True)
for source in SOURCES:
    payload = extract_workbook(source)
    (OUT / f"{source.stem}.json").write_text(json.dumps(payload, indent=2), encoding="utf-8")

(OUT / f"{DOC.stem}.json").write_text(json.dumps(extract_doc(DOC), indent=2), encoding="utf-8")
